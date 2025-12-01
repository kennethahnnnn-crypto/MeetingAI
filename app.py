import os
import time
import json
from io import BytesIO
from flask import Flask, render_template, request, jsonify, send_file
from werkzeug.utils import secure_filename
import google.generativeai as genai
from docx import Document

app = Flask(__name__)
app.config['JSON_AS_ASCII'] = False
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024 # 100MB Limit

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# --- CONFIGURE AI ---
API_KEY = os.environ.get("GEMINI_API_KEY", "AIzaSyDOa786xLCS9rSG1MKTdqp8gFiUKeYmxu0")
genai.configure(api_key=API_KEY)

# Use 1.5 Flash for stability with long audio
MODEL_NAME = 'gemini-2.5-flash'

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/download_word', methods=['POST'])
def download_word():
    try:
        data = request.json
        doc = Document()
        
        doc.add_heading(data.get('title', 'Meeting Minutes'), 0)
        
        doc.add_heading('📌 요약 (Summary)', level=1)
        doc.add_paragraph(data.get('summary', ''))
        
        doc.add_heading('⚡️ 액션 아이템 (Action Items)', level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '담당자'
        hdr_cells[1].text = '할 일'
        hdr_cells[2].text = '마감일'
        
        for item in data.get('action_items', []):
            row_cells = table.add_row().cells
            row_cells[0].text = item.get('owner', '-')
            row_cells[1].text = item.get('task', '-')
            row_cells[2].text = item.get('deadline', '-')

        doc.add_heading('⚖️ 결정 사항 (Decisions)', level=1)
        for decision in data.get('key_decisions', []):
            doc.add_paragraph(decision, style='List Bullet')

        doc.add_heading('⏱️ 타임라인 (Timeline)', level=1)
        for t in data.get('timeline', []):
            doc.add_paragraph(f"[{t.get('time')}] {t.get('topic')}")

        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return send_file(
            file_stream,
            as_attachment=True,
            download_name='meeting_minutes.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'audio' not in request.files: return jsonify({"error": "No file part"}), 400
    file = request.files['audio']
    if file.filename == '': return jsonify({"error": "No selected file"}), 400

    if file:
        filepath = ""
        audio_file = None
        try:
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            print(f"🎧 Uploading {filename}...")
            audio_file = genai.upload_file(path=filepath)
            
            # Wait for processing
            while audio_file.state.name == "PROCESSING":
                time.sleep(1)
                audio_file = genai.get_file(audio_file.name)

            if audio_file.state.name == "FAILED": raise ValueError("Audio processing failed.")

            print("✅ Ready. Analyzing...")
            
            prompt = """
            You are a professional Executive Assistant. Listen to this recording.
            
            CRITICAL RULES:
            1. LANGUAGE: ALL OUTPUT MUST BE IN KOREAN (한국어).
            2. HONESTY: Summarize EXACTLY what is said. Do not invent business topics.
            
            OUTPUT FORMAT (JSON ONLY):
            {
                "title": "Meeting Title (Korean)",
                "summary": "3-sentence summary (Polite Korean)",
                "action_items": [ {"owner": "Name", "task": "Task", "deadline": "Date"} ],
                "key_decisions": [ "Decision 1", "Decision 2" ],
                "timeline": [ {"time": "00:00", "topic": "Topic"} ]
            }
            """
            model = genai.GenerativeModel(MODEL_NAME)
            
            # Retry Logic
            try:
                response = model.generate_content([prompt, audio_file])
            except Exception as e:
                if "429" in str(e):
                    print("⚠️ Quota hit. Waiting 30s...")
                    time.sleep(30)
                    response = model.generate_content([prompt, audio_file])
                else: raise e

            clean_json = response.text.replace('```json', '').replace('```', '').strip()
            start = clean_json.find('{')
            end = clean_json.rfind('}') + 1
            return clean_json[start:end]

        except Exception as e:
            print(f"Error: {e}")
            return jsonify({"error": f"분석 실패: {str(e)}"}), 500
        
        finally:
            if filepath and os.path.exists(filepath): os.remove(filepath)
            if audio_file: 
                try: audio_file.delete()
                except: pass

if __name__ == '__main__':
    app.run(debug=True, port=5004)